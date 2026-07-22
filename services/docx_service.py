from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Mm
from docxtpl import DocxTemplate, InlineImage


STATUS_OPTIONS = ["In Progress", "Completed", "Delayed", "On Hold"]


class DocxService:
    def __init__(self, template_path: Path, output_dir: Path) -> None:
        self.template_path = template_path
        self.output_dir = output_dir

    def render_report(
        self,
        job_id: str,
        report_data: dict[str, Any],
        image_paths: list[Path],
    ) -> Path:
        doc = DocxTemplate(self.template_path)
        context = self._normalize_context(report_data)

        for idx in range(1, 5):
            if idx <= len(image_paths):
                context[f"photo{idx}"] = InlineImage(doc, str(image_paths[idx - 1]), width=Mm(145))
                context[f"photo{idx}_description"] = ""
            else:
                context[f"photo{idx}"] = ""
                context[f"photo{idx}_description"] = ""

        if len(image_paths) > 4:
            print(f"Warning: job {job_id} received {len(image_paths)} images; only first 4 rendered.")

        output_path = self.output_dir / f"daily_progress_report_{job_id}.docx"
        doc.render(context)
        doc.save(output_path)
        self._remove_unused_photo_slots(output_path, min(len(image_paths), 4))
        return output_path

    def _normalize_context(self, data: dict[str, Any]) -> dict[str, Any]:
        context = {key: self._blank_if_none(value) for key, value in data.items()}
        if not context.get("prepared_by") or context.get("prepared_by") == "Krishna Shelar":
            context["prepared_by"] = "Virendra Ghule"
        for key in ["work_summary", "resources_used", "issues", "next_day_plan"]:
            context[key] = [self._normalize_row(row) for row in data.get(key, [])]
        context["overall_status_line"] = self._format_status(data.get("overall_status"))
        return context

    def _normalize_row(self, row: dict[str, Any]) -> dict[str, str]:
        return {key: self._blank_if_none(value) for key, value in row.items()}

    def _blank_if_none(self, value: Any) -> Any:
        return "" if value is None else value

    def _format_status(self, selected: str | None) -> str:
        selected = (selected or "").strip()
        return "    ".join(
            f"[{'x' if option == selected else ' '}] {option}" for option in STATUS_OPTIONS
        )

    def _remove_unused_photo_slots(self, report_path: Path, rendered_photo_count: int) -> None:
        document = Document(report_path)
        rendered_photo_count = max(0, min(rendered_photo_count, 4))

        for paragraph in list(document.paragraphs):
            text = paragraph.text.strip()
            for idx in range(rendered_photo_count + 1, 5):
                if text.endswith(f"Photo {idx}"):
                    self._delete_element(paragraph._element)
                    break

        photo_section_index = None
        for idx, table in enumerate(document.tables):
            table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
            if "SITE PHOTOGRAPHS" in table_text:
                photo_section_index = idx
                break

        if photo_section_index is not None:
            photo_tables = document.tables[photo_section_index + 1 : photo_section_index + 5]
            for table in photo_tables[rendered_photo_count:]:
                self._delete_element(table._element)

        document.save(report_path)

    def _delete_element(self, element: Any) -> None:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
